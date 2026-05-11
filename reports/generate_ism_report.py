import json
import math
import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports"
ASSET_DIR = OUT_DIR / "assets"
DOCX_PATH = OUT_DIR / "RecruitAI_ISM_Report.docx"
DEMO_PATH = ROOT / "data" / "demo_seed.json"

TEAL = "0F766E"
DARK = "1E293B"
MUTED = "64748B"
LIGHT = "E2E8F0"
PALE = "F8FAFC"
GREEN = "16A34A"
ORANGE = "F59E0B"
RED = "DC2626"
BLUE = "2563EB"


def font(size=24, bold=False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for path in candidates:
        if os.path.exists(path):
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def hex_to_rgb(value):
    value = value.strip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def wrap_text(draw, text, fnt, max_width):
    words = str(text).split()
    if not words:
        return [""]
    lines = []
    current = words[0]
    for word in words[1:]:
        test = f"{current} {word}"
        if draw.textbbox((0, 0), test, font=fnt)[2] <= max_width:
            current = test
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def rounded_box(draw, xy, fill, outline="#CBD5E1", radius=20, width=3):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def center_text(draw, box, text, fnt, fill="#0F172A", line_gap=6):
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, fnt, x2 - x1 - 28)
    heights = [draw.textbbox((0, 0), line, font=fnt)[3] for line in lines]
    total_h = sum(heights) + line_gap * (len(lines) - 1)
    y = y1 + ((y2 - y1) - total_h) / 2
    for line, h in zip(lines, heights):
        w = draw.textbbox((0, 0), line, font=fnt)[2]
        draw.text((x1 + ((x2 - x1) - w) / 2, y), line, font=fnt, fill=fill)
        y += h + line_gap


def arrow(draw, start, end, fill="#334155", width=4):
    draw.line([start, end], fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 16
    p1 = (end[0] - size * math.cos(angle - math.pi / 7), end[1] - size * math.sin(angle - math.pi / 7))
    p2 = (end[0] - size * math.cos(angle + math.pi / 7), end[1] - size * math.sin(angle + math.pi / 7))
    draw.polygon([end, p1, p2], fill=fill)


def draw_actor(draw, cx, cy, label):
    draw.ellipse((cx - 18, cy - 62, cx + 18, cy - 26), outline=hex_to_rgb(DARK), width=4)
    draw.line((cx, cy - 26, cx, cy + 42), fill=hex_to_rgb(DARK), width=4)
    draw.line((cx - 42, cy, cx + 42, cy), fill=hex_to_rgb(DARK), width=4)
    draw.line((cx, cy + 42, cx - 36, cy + 92), fill=hex_to_rgb(DARK), width=4)
    draw.line((cx, cy + 42, cx + 36, cy + 92), fill=hex_to_rgb(DARK), width=4)
    center_text(draw, (cx - 100, cy + 108, cx + 100, cy + 166), label, font(22, True), fill="#0F172A")


def save_use_case_diagram(path):
    img = Image.new("RGB", (1600, 1000), "white")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1600, 1000), fill="#F8FAFC")
    draw.text((60, 40), "Use Case Diagram - RecruitAI Decision Support System", font=font(42, True), fill=hex_to_rgb(DARK))
    rounded_box(draw, (360, 140, 1240, 900), fill="#FFFFFF", outline="#94A3B8", radius=34, width=4)
    draw.text((405, 165), "System Boundary: RecruitAI", font=font(28, True), fill=hex_to_rgb(TEAL))

    draw_actor(draw, 170, 430, "HR / Recruiter")
    draw_actor(draw, 1430, 430, "Admin")
    draw_actor(draw, 170, 730, "Instructor / Viewer")

    use_cases = [
        ((470, 240, 760, 320), "Authenticate and manage account"),
        ((840, 240, 1130, 320), "Maintain master data"),
        ((470, 370, 760, 450), "Select job description"),
        ((840, 370, 1130, 450), "Rank candidates with AI scoring"),
        ((470, 500, 760, 580), "Review candidate profile and JD"),
        ((840, 500, 1130, 580), "Record Shortlist / Hold / Reject"),
        ((470, 630, 760, 710), "Track decision history"),
        ((840, 630, 1130, 710), "Monitor dashboard KPIs"),
        ((655, 770, 945, 850), "Deploy and run demo system"),
    ]
    for box, label in use_cases:
        rounded_box(draw, box, fill="#ECFEFF", outline="#0F766E", radius=42, width=3)
        center_text(draw, box, label, font(21, True), fill="#134E4A")

    hr_points = [(360, 275), (360, 410), (360, 540), (360, 670), (360, 810)]
    hr_targets = [(470, 280), (470, 410), (470, 540), (470, 670), (655, 810)]
    for start, end in zip(hr_points, hr_targets):
        arrow(draw, start, end, fill="#475569", width=3)
    for y, target in [(275, (1130, 280)), (410, (840, 410)), (540, (1130, 540)), (670, (1130, 670))]:
        arrow(draw, (1240, y), target, fill="#475569", width=3)
    arrow(draw, (260, 730), (470, 670), fill="#475569", width=3)
    arrow(draw, (1320, 730), (945, 810), fill="#475569", width=3)

    img.save(path)


def save_workflow_diagram(path):
    img = Image.new("RGB", (1700, 1100), "white")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1700, 1100), fill="#F8FAFC")
    draw.text((60, 40), "Workflow Diagram - Recruitment Decision Process", font=font(42, True), fill=hex_to_rgb(DARK))

    lanes = [
        (70, 150, 1630, 375, "HR User"),
        (70, 395, 1630, 720, "RecruitAI DSS"),
        (70, 740, 1630, 990, "Database and Decision Records"),
    ]
    for x1, y1, x2, y2, label in lanes:
        rounded_box(draw, (x1, y1, x2, y2), fill="#FFFFFF", outline="#CBD5E1", radius=24, width=2)
        draw.text((x1 + 25, y1 + 18), label, font=font(26, True), fill=hex_to_rgb(TEAL))

    steps = [
        ((180, 230, 390, 315), "Login"),
        ((500, 230, 730, 315), "Choose JD"),
        ((835, 230, 1075, 315), "Review ranked list"),
        ((1190, 230, 1460, 315), "Open profile + JD"),
        ((1190, 535, 1460, 625), "Decide: Shortlist / Hold / Reject"),
        ((835, 535, 1075, 625), "Compute score breakdown"),
        ((500, 535, 730, 625), "Batch embedding + skill match"),
        ((180, 535, 390, 625), "Fetch jobs and candidates"),
        ((500, 820, 730, 910), "Persist match scores"),
        ((835, 820, 1075, 910), "Save recruiter action"),
        ((1190, 820, 1460, 910), "Update dashboard / history"),
    ]
    for box, label in steps:
        fill = "#ECFEFF" if box[1] < 395 else "#EFF6FF" if box[1] < 740 else "#F0FDF4"
        outline = "#0F766E" if box[1] < 395 else "#2563EB" if box[1] < 740 else "#16A34A"
        rounded_box(draw, box, fill=fill, outline=outline, radius=22, width=3)
        center_text(draw, box, label, font(22, True), fill="#0F172A")

    arrows = [
        ((390, 272), (500, 272)),
        ((730, 272), (835, 272)),
        ((1075, 272), (1190, 272)),
        ((1325, 315), (1325, 535)),
        ((1190, 580), (1075, 580)),
        ((835, 580), (730, 580)),
        ((500, 580), (390, 580)),
        ((285, 625), (600, 820)),
        ((730, 865), (835, 865)),
        ((1075, 865), (1190, 865)),
        ((1325, 820), (1325, 625)),
    ]
    for start, end in arrows:
        arrow(draw, start, end, fill="#334155", width=4)

    draw.text((180, 1015), "Control points: authentication, role-based CRUD, cached AI features, decision audit trail, dashboard feedback loop.", font=font(24), fill=hex_to_rgb(MUTED))
    img.save(path)


def save_architecture_diagram(path):
    img = Image.new("RGB", (1700, 1050), "white")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1700, 1050), fill="#F8FAFC")
    draw.text((60, 40), "System Architecture Diagram - RecruitAI", font=font(42, True), fill=hex_to_rgb(DARK))

    layers = [
        (80, 150, 480, 930, "Presentation Layer", "#ECFEFF", "#0F766E"),
        (540, 150, 980, 930, "Application/API Layer", "#EFF6FF", "#2563EB"),
        (1040, 150, 1620, 930, "Data and AI Layer", "#F0FDF4", "#16A34A"),
    ]
    for x1, y1, x2, y2, title, fill, outline in layers:
        rounded_box(draw, (x1, y1, x2, y2), fill="#FFFFFF", outline="#CBD5E1", radius=28, width=2)
        draw.text((x1 + 24, y1 + 20), title, font=font(27, True), fill=hex_to_rgb(outline.strip("#")))

    boxes = [
        ((135, 250, 425, 350), "React + Vite SPA", "#ECFEFF", "#0F766E"),
        ((135, 420, 425, 540), "Routes: Dashboard, Jobs, Ranking, Detail, History, Master Data, Account", "#ECFEFF", "#0F766E"),
        ((135, 620, 425, 735), "React Query + Zustand state", "#ECFEFF", "#0F766E"),
        ((590, 250, 930, 350), "FastAPI REST endpoints under /api", "#EFF6FF", "#2563EB"),
        ((590, 420, 930, 540), "Auth: token, password hashing, account settings", "#EFF6FF", "#2563EB"),
        ((590, 620, 930, 735), "Recruitment workflow services and serializers", "#EFF6FF", "#2563EB"),
        ((1095, 230, 1390, 330), "SQLite operational database", "#F0FDF4", "#16A34A"),
        ((1095, 380, 1390, 505), "Sentence Transformers all-MiniLM-L6-v2", "#F0FDF4", "#16A34A"),
        ((1095, 555, 1390, 680), "CandidateJobMatcher: semantic, skill, experience, location scores", "#F0FDF4", "#16A34A"),
        ((1095, 730, 1390, 850), "Persistent candidate feature cache", "#F0FDF4", "#16A34A"),
        ((1415, 380, 1580, 505), "Docker Space on Hugging Face", "#FEF3C7", "#D97706"),
    ]
    for box, label, fill, outline in boxes:
        rounded_box(draw, box, fill=fill, outline=outline, radius=20, width=3)
        center_text(draw, box, label, font(21, True), fill="#0F172A")

    for start, end in [
        ((425, 300), (590, 300)),
        ((425, 480), (590, 480)),
        ((930, 300), (1095, 280)),
        ((930, 480), (1095, 442)),
        ((930, 665), (1095, 615)),
        ((1245, 330), (1245, 380)),
        ((1245, 505), (1245, 555)),
        ((1245, 680), (1245, 730)),
        ((1390, 442), (1415, 442)),
    ]:
        arrow(draw, start, end, fill="#334155", width=4)

    draw.text((120, 970), "Deployment model: one Docker container serves the React build and FastAPI API; demo data is seeded from JSON when SQLite DB is absent.", font=font(24), fill=hex_to_rgb(MUTED))
    img.save(path)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=DARK):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color="FFFFFF")
        set_cell_shading(hdr[i], TEAL)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if i == 0:
                set_cell_shading(cells[i], "F1F5F9")
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(TEAL if level == 1 else DARK)
        run.font.name = "Arial"
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.08
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor.from_string(DARK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(DARK)
    return p


def add_callout(doc, title, body, fill="ECFEFF"):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(TEAL)
    run.font.size = Pt(10.5)
    p.add_run("\n" + body)
    for r in p.runs[1:]:
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor.from_string(DARK)
    doc.add_paragraph()


def add_figure(doc, image_path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.55))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(MUTED)


def build_report():
    OUT_DIR.mkdir(exist_ok=True)
    ASSET_DIR.mkdir(exist_ok=True)
    use_case = ASSET_DIR / "use_case_diagram.png"
    workflow = ASSET_DIR / "workflow_diagram.png"
    architecture = ASSET_DIR / "system_architecture_diagram.png"
    save_use_case_diagram(use_case)
    save_workflow_diagram(workflow)
    save_architecture_diagram(architecture)

    payload = json.loads(DEMO_PATH.read_text(encoding="utf-8"))
    job_count = len(payload.get("jobs", []))
    candidate_count = len(payload.get("candidates", []))

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.65)
    section.bottom_margin = Cm(1.65)
    section.left_margin = Cm(1.85)
    section.right_margin = Cm(1.85)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = RGBColor.from_string(DARK)

    # Cover page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(90)
    r = p.add_run("RecruitAI")
    r.bold = True
    r.font.size = Pt(34)
    r.font.color.rgb = RGBColor.from_string(TEAL)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Information System Management Report")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor.from_string(DARK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Recruitment Decision Support System using semantic AI matching").italic = True
    add_callout(
        doc,
        "Prepared for: Information System Management",
        "Scope: use case analysis, workflow analysis, system architecture, data governance, security, deployment, and managerial value.",
        fill="F0FDFA",
    )
    add_table(
        doc,
        ["Project", "Technology Stack", "Demo Scale", "Default User"],
        [["RecruitAI", "React + Vite, FastAPI, SQLite, SQLAlchemy, Sentence Transformers", f"{job_count} jobs / {candidate_count} candidates", "HR/Admin"]],
        [Cm(3.0), Cm(6.7), Cm(3.2), Cm(2.7)],
    )
    doc.add_page_break()

    add_heading(doc, "Executive Summary", 1)
    add_body(
        doc,
        "RecruitAI is a recruitment decision support system designed to help HR users convert a large candidate pool into an explainable, auditable shortlist. "
        "The system combines master data management, semantic candidate ranking, detailed profile review, recruiter decisions, dashboard monitoring, and deployment automation."
    )
    add_body(
        doc,
        "From an Information System Management perspective, the project demonstrates how people, process, data, and technology can be integrated into one operational decision workflow. "
        "The system is not positioned as a fully autonomous hiring tool; instead, it augments HR judgment with transparent scoring components and preserves a human-in-the-loop decision record."
    )
    add_callout(
        doc,
        "Management value",
        "RecruitAI reduces screening time, standardizes evaluation evidence, improves traceability, and gives stakeholders a live view of recruitment workload and decisions.",
    )

    add_heading(doc, "1. Business Context and Problem Statement", 1)
    add_body(
        doc,
        "Recruitment teams often handle hundreds or thousands of profiles per job opening. Manual screening is slow, inconsistent, and difficult to audit. "
        "The main information management challenge is transforming unstructured job descriptions and candidate profiles into useful, comparable decision information."
    )
    add_bullet(doc, "Business problem: HR needs to find relevant candidates quickly without losing explainability.")
    add_bullet(doc, "Information problem: candidate skills, experience, desired role, and location are spread across semi-structured text fields.")
    add_bullet(doc, "Control problem: recruiter decisions must be stored with context so future review and reporting are possible.")

    add_heading(doc, "2. Project Scope", 1)
    add_table(
        doc,
        ["Scope Area", "Implemented Capability", "ISM Relevance"],
        [
            ["Authentication", "Login, token session, password change", "Access control and accountability"],
            ["Master Data", "Full CRUD for jobs, candidates, and decisions", "Data stewardship and operational ownership"],
            ["Decision Support", "AI ranking and score breakdown per candidate/job", "Decision quality and information processing"],
            ["Workflow", "Job selection -> ranking -> candidate detail -> decision", "Process standardization"],
            ["Monitoring", "Dashboard KPIs and decision history", "Management control and feedback"],
            ["Deployment", "Docker Space with seeded demo data", "System availability for stakeholders"],
        ],
        [Cm(3.4), Cm(6.0), Cm(6.2)],
    )

    add_heading(doc, "3. Stakeholders and Use Case Diagram", 1)
    add_body(
        doc,
        "The primary actor is the HR user or recruiter. The admin role maintains the master data and ensures the system remains usable. "
        "An instructor or viewer interacts mainly with the deployed demo, dashboard, and decision evidence during evaluation."
    )
    add_figure(doc, use_case, "Figure 1. Use case diagram for RecruitAI.")
    add_table(
        doc,
        ["Use Case", "Actor", "Outcome"],
        [
            ["Authenticate and manage account", "HR/Admin", "Only authorized users can use protected workflows."],
            ["Maintain master data", "Admin", "Jobs and candidate records remain accurate and searchable."],
            ["Select job description", "HR", "A specific JD becomes the basis for matching."],
            ["Rank candidates", "HR", "System returns top candidates with score breakdowns."],
            ["Review candidate profile and JD", "HR", "Decision maker sees evidence before acting."],
            ["Record decision", "HR", "Shortlist, Hold, or Reject is stored for audit/history."],
            ["Monitor dashboard", "HR/Admin", "Recruitment status and quality indicators are visible."],
        ],
        [Cm(5.1), Cm(3.0), Cm(7.4)],
    )

    add_heading(doc, "4. Workflow Analysis", 1)
    add_body(
        doc,
        "The workflow is organized around a human-in-the-loop decision cycle. The system recommends and explains; the recruiter decides. "
        "This structure is appropriate for a management information system because it improves decision efficiency while preserving managerial control."
    )
    add_figure(doc, workflow, "Figure 2. End-to-end recruitment workflow.")
    add_table(
        doc,
        ["Workflow Step", "System Support", "Control / Output"],
        [
            ["Login", "Token-based authenticated session", "User accountability"],
            ["Select JD", "Job catalog and selected job state", "Clear decision context"],
            ["AI ranking", "Semantic matching, skill overlap, experience/location scoring", "Candidate list ranked by evidence"],
            ["Detail review", "Profile, JD, score details, matched/missing skills", "Transparent recommendation"],
            ["Decision", "Shortlist/Hold/Reject and notes", "Audit-ready recruiter action"],
            ["Feedback", "Dashboard and history update", "Management reporting loop"],
        ],
        [Cm(3.2), Cm(7.0), Cm(5.2)],
    )

    add_heading(doc, "5. System Architecture", 1)
    add_body(
        doc,
        "RecruitAI follows a three-layer architecture: a React single-page application, a FastAPI application service, and a data/AI layer. "
        "For deployment, the React build is served by FastAPI inside one Docker container, which simplifies access for classroom demonstration."
    )
    add_figure(doc, architecture, "Figure 3. System architecture diagram.")
    add_table(
        doc,
        ["Layer", "Components", "Responsibility"],
        [
            ["Presentation", "React, Vite, React Query, Zustand, Tailwind UI", "Navigation, data forms, ranking table, dashboard, decision screens"],
            ["Application/API", "FastAPI, Pydantic models, auth helpers, serializers", "Expose REST APIs, validate payloads, enforce authentication, orchestrate workflow"],
            ["Data", "SQLite, SQLAlchemy models", "Persist jobs, candidates, matches, recruiter actions, users"],
            ["AI/Decision", "Sentence Transformers, CandidateJobMatcher, feature cache", "Generate embeddings, compute score components, rank candidates"],
            ["Deployment", "Dockerfile, HF Space startup script, demo seed JSON", "Package and run the full system through one public link"],
        ],
        [Cm(3.1), Cm(5.6), Cm(6.8)],
    )

    add_heading(doc, "6. Data Management Design", 1)
    add_body(
        doc,
        "The operational database stores five central entities: Job, Candidate, Match, RecruiterAction, and User. "
        "This model separates source data from decision records, which is important because the same candidate may be evaluated against multiple jobs."
    )
    add_table(
        doc,
        ["Entity", "Key Fields", "Management Purpose"],
        [
            ["Job", "job_id, title, description, requirements, location, salary", "Represents the demand side of recruitment."],
            ["Candidate", "user_id, name, desired_job, skills, experience, degree, location", "Represents the applicant or talent pool."],
            ["Match", "job_id, user_id, fit_score, component scores", "Stores computed decision-support evidence."],
            ["RecruiterAction", "decision, notes, recruiter_name, timestamps", "Stores final human action and audit trail."],
            ["User", "username, password_hash, role", "Controls access to protected functions."],
        ],
        [Cm(3.2), Cm(6.2), Cm(6.0)],
    )

    add_heading(doc, "7. Decision Support Logic", 1)
    add_body(
        doc,
        "The matching engine combines semantic text similarity with structured business rules. Candidate and job text are embedded with all-MiniLM-L6-v2, producing 384-dimensional vectors. "
        "The final score is then computed from semantic similarity, skill overlap, experience match, and location match."
    )
    add_table(
        doc,
        ["Score Component", "Default Weight", "How It Supports HR Decision"],
        [
            ["Text similarity", "35%", "Captures semantic fit between JD and candidate desired role/skills."],
            ["Skill match", "35%", "Measures overlap between extracted job requirements and candidate skills."],
            ["Experience match", "20%", "Rewards candidates meeting the required years of experience."],
            ["Location match", "10%", "Adds operational fit and applies penalty for explicit mismatch."],
        ],
        [Cm(4.2), Cm(2.8), Cm(8.5)],
    )
    add_callout(
        doc,
        "Accuracy-preserving optimization",
        "The project caches candidate embeddings and extracted skills by content hash. This reduces ranking latency without changing the scoring formula or model.",
        fill="EFF6FF",
    )

    add_heading(doc, "8. Information System Management Evaluation", 1)
    add_table(
        doc,
        ["ISM Dimension", "Assessment", "Project Evidence"],
        [
            ["People", "Supports HR staff rather than replacing them", "Human recruiter still records final decision."],
            ["Process", "Standardizes screening workflow", "Job selection, ranking, profile review, decision history."],
            ["Data", "Turns semi-structured text into managed decision data", "SQLite schema plus demo seed data."],
            ["Technology", "Uses AI/NLP within a web information system", "FastAPI, React, SQLAlchemy, sentence-transformers."],
            ["Control", "Provides authentication and audit trail", "Token auth, password change, recruiter actions."],
            ["Performance", "Improves responsiveness for large candidate lists", "Cached features and paginated master data."],
        ],
        [Cm(3.2), Cm(6.1), Cm(6.1)],
    )
    add_body(
        doc,
        "The system's management value is strongest in the middle layer between raw data and final hiring decision. It does not merely store records; it processes information into ranked alternatives, "
        "explanatory evidence, and dashboard feedback. This aligns with the purpose of a decision support system in management information systems."
    )

    add_heading(doc, "9. Security, Privacy, and Governance", 1)
    add_table(
        doc,
        ["Risk", "Current Control", "Recommended Improvement"],
        [
            ["Unauthorized data changes", "Authenticated CRUD endpoints", "Add role granularity and server-side authorization policies."],
            ["Weak default credentials", "Password change screen exists", "Force password change on first login in production."],
            ["Candidate privacy exposure", "Local SQLite and demo deployment", "Mask sensitive fields and define data retention policy."],
            ["Algorithmic bias", "Human-in-the-loop decision", "Add fairness testing across gender/location/experience segments."],
            ["Ephemeral free hosting storage", "Seed demo data on startup", "Use persistent storage for production use."],
        ],
        [Cm(4.0), Cm(5.2), Cm(6.2)],
    )

    add_heading(doc, "10. Deployment and Operations", 1)
    add_body(
        doc,
        "For classroom demonstration, the project is packaged as a Hugging Face Docker Space. The Docker image builds the frontend, installs backend dependencies, seeds the database from JSON when needed, "
        "and runs FastAPI on the required web port. This creates a single public URL for the instructor."
    )
    add_table(
        doc,
        ["Operational Concern", "Implementation"],
        [
            ["Startup", "scripts/start_hf_space.py sets environment paths and runs seed.py if data/app.db is missing."],
            ["Data bootstrap", f"data/demo_seed.json contains {job_count} jobs and {candidate_count} candidates."],
            ["Model cache", "HF_HOME and TRANSFORMERS_CACHE point to data/processed/hf_cache."],
            ["Frontend delivery", "FastAPI serves frontend/dist and /assets in the same container."],
            ["Known limitation", "Free Hugging Face storage is ephemeral, so demo edits may reset after restart."],
        ],
        [Cm(4.6), Cm(10.8)],
    )

    add_heading(doc, "11. Limitations and Future Improvements", 1)
    add_bullet(doc, "Add role-based access control for separate HR, Admin, and Viewer permissions.")
    add_bullet(doc, "Add structured import/export pipelines for CSV/Excel master data maintenance.")
    add_bullet(doc, "Add fairness dashboards and bias checks before using the tool for real hiring decisions.")
    add_bullet(doc, "Add persistent cloud database for production instead of ephemeral demo SQLite.")
    add_bullet(doc, "Add interview scheduling and email notification integrations to complete the recruitment lifecycle.")

    add_heading(doc, "12. Conclusion", 1)
    add_body(
        doc,
        "RecruitAI is a strong Information System Management project because it demonstrates a complete information lifecycle: data capture, processing, decision support, human action, storage, reporting, and deployment. "
        "The system creates practical value for HR by reducing screening friction while maintaining transparency and accountability."
    )

    add_heading(doc, "References", 1)
    refs = [
        "Project source code: Recruitment Decision Support System repository.",
        "Backend framework: FastAPI application source in backend/main.py.",
        "Frontend framework: React/Vite application source in frontend/src.",
        "Data model: SQLAlchemy ORM models in backend/src/storage/models.py.",
        "AI model: Sentence Transformers all-MiniLM-L6-v2 usage in backend/src/models/embedder.py.",
        "Deployment: Dockerfile and Hugging Face Spaces startup script in scripts/start_hf_space.py.",
    ]
    for ref in refs:
        add_bullet(doc, ref)

    # Footer
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "RecruitAI - Information System Management Report"
        for run in footer.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string(MUTED)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_report()
